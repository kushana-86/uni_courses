from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("xiaomimimo_application_zhouxingyu_fixed.docx")
IMG_DIR = Path("outputs")


def font_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font_run(r, 13 if level == 1 else 11.5, True)
    return p


def para(doc, text="", bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    font_run(r, 10.5, bold)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    font_run(r, 10.5)
    return p


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


q04 = (
    "我构建的核心 AI 项目是“志摩凛 AI 便携陪伴 Agent”，并围绕它扩展了多个检索、分析和小程序项目。"
    "志摩凛 AI 解决的痛点是：普通聊天机器人缺少长期记忆、角色一致性、语音/视觉陪伴和离线可用性，难以真正成为随身设备上的个人 Agent。"
    "我用 Node.js 搭建本地服务，前端提供 5-7 英寸小屏 UI、Live2D/Three.js 3D 模型、语音输入和语音播报；"
    "后端包含 persona prompt、记忆管理、Skill 路由、闹钟工具、远端模型调用和本地规则回退。"
    "核心逻辑流是：用户文字或录音输入后，系统先识别语言与意图，再匹配露营规划、治愈陪伴、户外知识、时间/闹钟等 Skill；"
    "随后从本地 user-profile 取出姓名、偏好、计划、重要日期和聊天风格，构建角色系统提示词，优先调用 DeepSeek/兼容 OpenAI 接口生成回复，失败时自动回退本地规则引擎；"
    "最后触发浏览器 TTS、OpenVoice 或 GPT-SoVITS 声线链路，并更新 Live2D/3D 状态。"
    "项目还做了离线版 shima-rin-ai-offline，面向树莓派和旅行随身设备，支持日本时区、旅行清单、常用日语、离线 TTS、闹钟和记忆。"
    "除此之外，我还做了 INF_A 自动信息分析项目，用 Python 自动清洗国家统计局失业率数据、完成相关/偏相关/逐步回归和报告生成；"
    "做了“择域先知”微信小程序，帮助非西藏生源定向西藏就业用户检索、收藏和对比区县岗位；"
    "做了 civ6-city-decision-system，将游戏开城问题拆成知识检索、评分、策略推荐和问答解释。"
    "整体上，我的项目不是单次 prompt，而是持续构建“感知输入-检索记忆/知识-多工具推理-生成输出-语音/界面反馈-持久化记忆”的 Agent 工作流。"
)

q05_text = (
    "可提交志摩凛 AI 本地运行截图/录屏、终端日志、项目目录或 GitHub 仓库截图、语音合成文件和测试结果作为证明。"
    "优先展示：启动 npm start 后的小屏 UI、当前 Skill 切换、记忆面板更新、Live2D/3D 模型切换、闹钟工具、OpenVoice/GPT-SoVITS 语音链路，"
    "以及离线版在本地 TTS 模式下的运行状态。再补充 INF_A、择域先知微信小程序和 civ6 决策系统的仓库/页面截图，"
    "用来证明我有持续构建 AI/Agent 产品和落地工具的能力。"
)

projects = [
    (
        "志摩凛 AI 便携陪伴 Agent",
        "E:/shima-rin-ai；离线版 E:/shima-rin-ai-offline",
        "DeepSeek/兼容模型接入、本地规则回退、长期记忆、Skill 路由、闹钟工具、Live2D/3D、OpenVoice/GPT-SoVITS、离线 TTS。",
    ),
    (
        "AI/信息分析仓库 INF_A",
        "https://github.com/kushana-86/INF_A",
        "自动清洗统计数据，生成相关分析、偏相关、逐步回归、图表、Excel 与报告。",
    ),
    (
        "微信小程序 择域先知",
        "https://github.com/kushana-86/zyxz-tibet-job-helper-miniapp",
        "非西藏生源定向西藏就业选岗助手，支持搜索、收藏、区县对比和综合适宜度评分。",
    ),
    (
        "文明六决策系统",
        "https://github.com/kushana-86/civ6-city-decision-system",
        "游戏开城决策 Agent 原型，包含检索、评分、策略推荐、知识图谱和本地问答。",
    ),
    (
        "语义重排序实验",
        "本地目录：jiansuo/231001201_周星宇",
        "可复现 RAG/知识库检索排序实验，自动输出指标、图表和 Word 论文。",
    ),
]

materials = [
    "志摩凛 AI 运行截图/录屏：展示输入问题、触发 Skill、回复生成、记忆面板更新、Live2D/3D 状态变化。",
    "终端日志：在 E:/shima-rin-ai 或 E:/shima-rin-ai-offline 运行 npm start、npm test 的输出。",
    "语音链路证明：OpenVoice/GPT-SoVITS 或离线 TTS 生成音频的截图/录屏，或 .tmp-audio、sound、gpt-sovits 数据集目录截图。",
    "微信小程序截图：择域先知在微信开发者工具中的总览、搜索、对比、候选收藏页面。",
    "GitHub/结果证明：INF_A、zyxz-tibet-job-helper-miniapp、civ6-city-decision-system 的仓库页面，以及语义重排序指标图。",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("xiaomimimo 创造者 Token 激励计划 / Agent 生态共建计划申请材料")
    font_run(r, 16, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("周星宇 | 重点项目：志摩凛 AI 便携陪伴 Agent")
    font_run(r, 10, False, (90, 90, 90))

    heading(doc, "04 表单粘贴版", 1)
    para(doc, "请直接复制下面这段到第 04 题：", True)
    para(doc, q04)
    para(doc, f"字数提示：约 {len(q04)} 个中文字符，低于 1200 字上限。")

    heading(doc, "05 表单填写与上传建议", 1)
    para(doc, "第 05 题可填写/备注：", True)
    para(doc, q05_text)
    para(doc, "建议上传优先级：", True)
    for item in materials:
        bullet(doc, item)

    heading(doc, "项目与证明对照", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["项目", "链接/位置", "证明点"]):
        table.rows[0].cells[i].text = h
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in projects:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    heading(doc, "可量化亮点", 1)
    bullet(doc, "志摩凛 AI：已形成文字/录音输入、角色 prompt、记忆注入、Skill 路由、工具调用、语音播报、小屏 UI 和离线版部署的完整 Agent 链路。")
    bullet(doc, "语义重排序：Top-1 0.667→0.750，Hit@5 0.833→1.000，MRR 0.742→0.840，nDCG@5 0.761→0.880。")
    bullet(doc, "INF_A：35 个月国家统计局数据样本，自动完成相关分析、偏相关、距离分析、一元/多元回归，最优模型调整 R²=0.9450。")
    bullet(doc, "择域先知：围绕西藏区县选岗，沉淀本地数据、风险标签、交通/海拔/工作半径等多指标评分和候选管理闭环。")

    for img_name, caption in [("metric_chart.png", "语义重排序指标提升图"), ("architecture.png", "语义重排序流程图")]:
        img = IMG_DIR / img_name
        if img.exists():
            heading(doc, caption, 2)
            doc.add_picture(str(img), width=Cm(15))
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "提交提醒", 1)
    para(
        doc,
        "如果上传文件数有限，优先上传志摩凛 AI 录屏或截图、终端运行日志、语音链路截图；再上传微信小程序和 GitHub 仓库截图。"
        "这样第 04 题描述的 Agent 主项目与第 05 题证明材料能完全对应。",
    )

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
