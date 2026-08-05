from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(".")
OUT = ROOT / "outputs"
PAPER = ROOT / "语义重排序信息检索优化论文_扩展版.docx"


def set_style(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "黑体"


def h(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def p(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing = 1.35
    para.add_run(text)
    return para


def plain(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.25
    para.add_run(text)
    return para


def add_image(doc: Document, filename: str, caption: str, width: float = 15.5):
    path = OUT / filename
    if path.exists():
        doc.add_picture(str(path), width=Cm(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_metric_table(doc: Document):
    metrics = pd.read_csv(OUT / "metrics.csv", index_col=0)
    table_df = metrics.reset_index().rename(columns={"index": "方法"}).round(3)
    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(table_df.columns):
        table.rows[0].cells[i].text = str(col)
    for row in table_df.itertuples(index=False):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = f"{val:.3f}" if isinstance(val, float) else str(val)
    cap = doc.add_paragraph("表1  TF-IDF与语义重排序方法的检索指标对比")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return metrics


def add_rank_table(doc: Document):
    rank = pd.read_csv(OUT / "rank_change.csv")
    show = rank[["查询编号", "相关文档", "TF-IDF排名", "重排序排名", "TF-IDF Top1", "重排序Top1"]]
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    for i, col in enumerate(show.columns):
        table.rows[0].cells[i].text = str(col)
    for row in show.itertuples(index=False):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    cap = doc.add_paragraph("表2  查询级相关文档排名变化")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return rank


def build():
    doc = Document()
    set_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于语义重排序的信息检索结果优化方法研究")
    run.bold = True
    run.font.size = Pt(18)

    plain(doc, "摘要")
    p(
        doc,
        "信息检索系统是搜索引擎、数字图书馆、企业知识库和检索增强生成系统的基础组件。随着文本资源规模持续增长，用户对检索系统的要求已经从“能够找到相关内容”进一步转向“能够把最相关、最可信、最能回答问题的内容排在最前面”。传统TF-IDF、BM25等稀疏检索方法具有实现简单、速度快、可解释性强等优势，至今仍是工程系统中常用的召回基线；但这类方法主要依赖词面匹配，当用户查询与文档使用近义表达、抽象表达或上下文相关表达时，容易出现词汇不匹配问题。为改善这一问题，本文围绕“初始检索+语义重排序”的两阶段检索框架展开研究。首先使用字符n-gram TF-IDF完成候选文档召回，然后通过同义词归一、潜在语义分析和余弦相似度计算实现轻量级语义重排序。实验构建40条中文检索语料和12个查询样例，对比原始TF-IDF排序与语义重排序方法在Top-1准确率、Hit@K、MRR和nDCG等指标上的表现。结果显示，重排序后Top-1准确率由0.667提升至0.750，Hit@5由0.833提升至1.000，MRR由0.742提升至0.840，说明语义重排序能够有效改善靠前结果的相关性。本文进一步分析了方法有效原因、实验局限和在RAG问答场景中的应用价值，为小规模中文检索系统的排序优化提供了可复现参考。",
    )
    plain(doc, "关键词：信息检索；语义重排序；TF-IDF；潜在语义分析；RAG；排序优化")

    h(doc, "1 引言")
    h(doc, "1.1 研究背景", 2)
    p(
        doc,
        "信息检索是现代信息系统中最基础、也最容易被用户直接感知的能力之一。无论是通用搜索引擎、学术论文数据库、校园图书馆系统，还是企业内部知识库和智能问答平台，检索系统都承担着在大量文档中快速定位有用信息的任务。随着数字化办公、在线教育和生成式人工智能应用的普及，文本资源的规模和更新速度不断提高，用户查询也逐渐从简单关键词扩展为自然语言问题、任务式指令和多轮问答上下文。在这种背景下，检索系统的质量不再只由“是否能返回结果”决定，而更取决于“最有价值的结果是否足够靠前”。",
    )
    p(
        doc,
        "从用户体验角度看，排序质量直接影响信息获取效率。用户通常只浏览前几条或第一页结果，如果正确答案排在较后位置，即使系统实际上已经召回了相关文档，用户也可能认为检索失败。特别是在问答系统和RAG系统中，排序错误的影响会进一步放大：生成模型往往只读取排名靠前的若干个文本片段，如果这些片段与问题关系较弱，模型就可能生成不完整、不准确甚至带有幻觉的回答。因此，检索排序已经成为连接传统信息检索与大模型应用的重要环节。",
    )
    p(
        doc,
        "传统信息检索方法长期以词项匹配为核心。TF-IDF通过词频和逆文档频率计算词项权重，BM25在此基础上进一步考虑词频饱和和文档长度归一化。这些方法计算效率高、工程实现成熟，在大规模召回阶段仍具有不可替代的作用。然而，词项匹配方法天然依赖查询与文档之间的字面重合。当用户输入“怎样把搜索结果按含义做优化”时，真正相关文档可能使用“语义重排序”“候选文档”“余弦相似度”等表达；如果查询和文档缺少足够重合词，传统方法就可能低估其相关性。",
    )
    h(doc, "1.2 问题提出", 2)
    p(
        doc,
        "本文关注的问题可以概括为：在不显著增加系统复杂度的前提下，如何提升初始检索结果的语义相关性，使真正符合用户意图的文档在最终结果中尽可能靠前。这个问题并不等同于完全替代传统检索。相反，在真实工程系统中，更常见也更稳健的方案是保留稀疏检索作为高效召回模块，再在较小候选集上引入语义模型进行重排序。这样既能利用TF-IDF、BM25等方法的速度优势，又能利用语义表示缓解词汇不匹配。",
    )
    p(
        doc,
        "重排序方法的核心思想是“先粗后精”。初始召回阶段从全部文档中快速筛选出可能相关的候选；重排序阶段只处理候选集合，因此可以使用更复杂的特征、更高维的向量表示或更强的神经网络模型。对于教学实验和小型系统而言，完全下载并部署大型预训练模型并不总是现实，因此本文采用一种轻量级语义重排序方法：利用人工可解释的同义词归一缓解近义表达差异，再利用潜在语义分析把文本投影到低维语义空间，最后用余弦相似度重新排序。",
    )
    h(doc, "1.3 研究意义", 2)
    p(
        doc,
        "本文研究具有三方面意义。第一，在理论层面，它展示了从词面匹配到语义匹配的排序优化思路，能够帮助理解现代检索系统为什么通常采用多阶段架构。第二，在实践层面，本文给出了一个完全可复现的小型中文实验，不依赖外部在线模型和复杂部署环境，适合在课程论文、实验报告和小型知识库原型中复用。第三，在应用层面，本文方法与RAG系统的检索模块具有直接关联，可作为理解“检索质量影响生成质量”的基础实验。",
    )
    h(doc, "1.4 本文主要工作", 2)
    p(
        doc,
        "本文主要完成以下工作：一是梳理传统信息检索、语义表示、重排序和RAG相关研究，明确两阶段检索框架的适用场景；二是设计一种“字符n-gram TF-IDF召回+同义词归一+LSA语义重排序”的轻量级方法；三是构建中文短文本实验数据集，设置12个查询和标准相关文档；四是从Top-1准确率、Hit@3、Hit@5、MRR、nDCG等多个指标进行对比实验；五是结合典型查询案例分析重排序有效与失效的原因，并讨论后续使用SBERT、BGE、Cross-Encoder或大语言模型排序器进行扩展的方向。",
    )

    h(doc, "2 相关技术与研究现状")
    h(doc, "2.1 传统词匹配检索模型", 2)
    p(
        doc,
        "传统信息检索模型主要建立在词项统计和概率排序思想之上。TF-IDF认为，如果一个词在某篇文档中出现频率较高，而在整个语料库中出现频率较低，则这个词对该文档具有较强区分能力。其基本权重可以表示为：w(t,d)=tf(t,d)×idf(t)，其中tf表示词项t在文档d中的出现频率，idf表示词项t在全体文档中的逆文档频率。该方法直观、简单、可解释性强，因此常被用于文本分类、关键词抽取和基础检索排序。",
    )
    p(
        doc,
        "BM25是经典概率检索框架中的代表方法，它在TF-IDF思想基础上加入词频饱和与文档长度归一化机制。与原始TF-IDF相比，BM25更适合处理长短不一的文档，也更符合搜索引擎排序的经验规律。尽管近年来稠密检索和大模型检索发展迅速，BM25仍然是许多公开评测和工业系统中的强基线。这说明词匹配方法并未过时，而是更适合作为多阶段检索中的第一阶段。",
    )
    p(
        doc,
        "词匹配模型的主要不足在于语义理解能力有限。它很难知道“搜索”“查找”“检索”在很多语境下表达相近含义，也难以理解“减少胡编”与“缓解幻觉”之间的语义联系。对于短查询而言，这一问题更加明显，因为短查询包含的词项少，任何一个关键表达不匹配都会显著影响得分。因此，仅依赖词面相似度往往无法满足问答式检索和知识库检索的需求。",
    )
    h(doc, "2.2 语义表示与向量检索", 2)
    p(
        doc,
        "语义表示方法试图把文本映射为连续向量，使语义相近的文本在向量空间中距离更近。早期方法包括潜在语义分析，它通过对词项-文档矩阵进行降维，发现词项之间的隐含关联。随着深度学习发展，Word2Vec、BERT、Sentence-BERT以及中文BGE等模型进一步提升了文本语义表示能力。Sentence-BERT通过孪生网络结构生成句向量，使句子相似度可以用余弦相似度高效计算；BGE等模型则针对检索任务进行训练，在中文和多语言场景中被广泛使用。",
    )
    p(
        doc,
        "向量检索通常也被称为稠密检索。它与TF-IDF、BM25等稀疏检索不同，不再只依赖离散词项是否重合，而是利用模型学习得到的连续向量表达语义关系。稠密检索能够缓解同义词、近义表达和上下文差异问题，但也存在部署成本高、模型更新复杂、可解释性弱等不足。因此，许多系统采用混合检索，即同时使用稀疏检索和稠密检索，再通过融合或重排序得到最终结果。",
    )
    h(doc, "2.3 重排序方法", 2)
    p(
        doc,
        "重排序是信息检索系统中承上启下的环节。初始召回阶段一般追求高召回率和低延迟，排序阶段则追求更高的相关性判断精度。常见重排序方法包括基于规则的重排序、基于特征融合的学习排序、基于双塔模型的向量重排序、基于Cross-Encoder的精排以及基于大语言模型的相关性判断。其中Cross-Encoder通常把查询和候选文档拼接后同时输入模型，能够捕捉细粒度交互信息，但计算成本较高，适合候选集较小的精排阶段。",
    )
    p(
        doc,
        "重排序的评价不应只关注是否命中相关文档，还应关注相关文档出现的位置。Top-1准确率反映第一条结果是否正确，Hit@K反映相关文档是否进入前K，MRR关注第一个相关结果的排名倒数，nDCG则考虑排名位置带来的折损。对于用户体验而言，把相关文档从第十位提升到第二位往往比单纯扩大候选集更有价值，因为用户实际浏览的范围有限。",
    )
    h(doc, "2.4 RAG场景下的检索排序", 2)
    p(
        doc,
        "检索增强生成把外部知识库与生成模型结合起来，常见流程包括文档切分、向量化、索引构建、查询召回、重排序和答案生成。RAG的目标是让模型在回答时参考可追溯的外部资料，从而缓解知识过时和幻觉问题。然而，如果检索阶段返回的上下文不相关，即使生成模型能力很强，也难以得到可靠答案。因此，RAG系统中的排序质量不仅影响检索结果本身，也会影响最终生成内容的准确性、完整性和可解释性。",
    )
    p(
        doc,
        "近年来关于RAG和大模型信息检索的研究表明，检索模块正在从单纯的文档查找演变为面向任务的知识选择。查询改写、混合检索、语义重排序、上下文压缩和答案一致性验证逐渐成为系统优化重点。本文虽然没有直接部署大型语言模型，但围绕语义重排序设计的实验能够体现RAG系统中的关键问题：候选文档并非只要被召回就足够，排序位置同样决定了后续模块能否获得正确依据。",
    )

    h(doc, "3 方法设计")
    h(doc, "3.1 总体框架", 2)
    add_image(doc, "architecture.png", "图1  初始检索与语义重排序流程图")
    p(
        doc,
        "本文设计的检索优化框架采用典型两阶段结构。第一阶段为初始检索模块，输入用户查询后使用TF-IDF计算查询与文档集合之间的词面相似度，得到候选文档列表。第二阶段为语义重排序模块，对查询和候选文档进行同义词归一与语义向量表示，再依据语义相似度重新调整候选文档顺序。最终输出的排序结果既保留了初始检索的效率，又引入了语义匹配能力。",
    )
    p(
        doc,
        "该框架的关键在于职责划分。初始检索阶段解决“从大量文档中快速找出可能相关结果”的问题，重排序阶段解决“在候选结果中判断谁更符合用户意图”的问题。两阶段结构降低了语义模型的计算压力，也更符合实际系统设计习惯。对于大规模语料，Top-K可以设置为几十到几百；对于本文40条文档的小规模实验，为观察完整候选列表的排序变化，实验中K取40。",
    )
    h(doc, "3.2 初始检索模块", 2)
    p(
        doc,
        "初始检索模块采用字符n-gram TF-IDF。中文文本不像英文文本那样天然以空格分隔，如果引入分词工具，实验复现会额外依赖词典和分词模型。为降低环境要求，本文使用字符级n-gram特征，n取2至4。该方法能够捕捉“信息检索”“语义重排”“向量检索”等连续片段，在小规模中文语料中具有较好的稳定性。查询和文档被转换为稀疏向量后，使用余弦相似度计算词面相关得分。",
    )
    p(
        doc,
        "余弦相似度衡量两个向量夹角的接近程度，公式为sim(q,d)=q·d/(||q||·||d||)。当查询和文档共享较多高权重n-gram时，余弦相似度较高。该方法对完全相同或高度重合的表达效果较好，但当用户使用抽象表达或近义说法时，得分可能偏低。例如“首页结果不相关为什么会让用户流失”与“排序质量直接影响用户体验”表达相关，但字面重合并不总是充分。",
    )
    h(doc, "3.3 语义重排序模块", 2)
    p(
        doc,
        "语义重排序模块首先进行同义词归一。实验中构建了若干领域同义词组，例如“搜索、查找、召回、检索”映射到相近语义，“重新排列、精排、排序优化、重排序”映射到排序优化语义，“生成答案、外部资料、减少胡编、幻觉”映射到RAG生成语义。这一步并不追求覆盖所有语言现象，而是通过可解释的方式缓解小型实验中的典型词汇不匹配问题。",
    )
    p(
        doc,
        "随后，系统将扩展后的查询和文档共同构造成TF-IDF矩阵，并使用TruncatedSVD进行潜在语义分析。LSA的思想是通过矩阵降维保留主要语义结构，减少词项噪声影响，使具有相似上下文的词和文档在低维空间中更加接近。降维后再进行向量归一化，避免文本长度差异对相似度造成过强影响。最后，在候选集中计算查询向量与文档向量的余弦相似度，并与初始词面得分进行加权融合。",
    )
    p(
        doc,
        "本文最终得分设置为Score=0.35×LexicalScore+0.65×SemanticScore。其中LexicalScore表示归一化后的TF-IDF词面得分，SemanticScore表示LSA语义向量相似度。权重设置体现了本文对语义匹配的强调，同时保留一定词面匹配约束，避免语义扩展过度导致无关文档被提升。该权重并非唯一最优选择，在更大数据集上可以通过验证集调参确定。",
    )
    h(doc, "3.4 算法流程", 2)
    for step in [
        "Step1：输入用户查询q和文档集合D。",
        "Step2：使用字符n-gram TF-IDF构建查询向量和文档向量。",
        "Step3：计算词面余弦相似度，得到初始候选文档列表。",
        "Step4：对查询和候选文档进行同义词归一与领域表达扩展。",
        "Step5：基于扩展文本构建TF-IDF矩阵，并通过TruncatedSVD得到低维语义向量。",
        "Step6：计算查询与候选文档的语义相似度。",
        "Step7：融合词面得分和语义得分，按照最终得分降序输出结果。",
    ]:
        plain(doc, step)
    p(
        doc,
        "从复杂度角度看，初始TF-IDF检索主要消耗在向量化和稀疏矩阵相似度计算上；语义重排序只在候选集合中进行，因此在候选集较小时计算成本可控。若后续替换为深度学习模型，仍可保持同样的两阶段框架，只需将LSA语义表示模块替换为句向量模型或Cross-Encoder评分模块。",
    )

    h(doc, "4 实验设计")
    h(doc, "4.1 实验环境", 2)
    p(
        doc,
        "实验在Python环境下完成，主要依赖sklearn、numpy、pandas、matplotlib和python-docx。为保证复现便利，实验没有下载外部预训练模型，也没有依赖中文分词工具。所有语料、查询、标准答案、实验指标和论文图表均由脚本自动生成。运行reproduce_semantic_rerank.py可以复现数据集、指标表、排名变化表、流程图和图片；运行generate_expanded_paper.py可以生成扩展版论文文档。",
    )
    h(doc, "4.2 数据集构建", 2)
    p(
        doc,
        "本文自构40条中文短文档，主题覆盖TF-IDF、BM25、语义向量、重排序、Cross-Encoder、RAG、评价指标、中文n-gram特征、企业知识库和排序用户体验等内容。每条文档均为一段概念性说明，长度较短，适合模拟知识库条目或教学资料片段。查询集包含12个中文自然语言查询，每个查询人工指定1条最相关文档作为标准答案。",
    )
    p(
        doc,
        "查询设计刻意加入近义表达和抽象表达，以测试词面匹配方法的不足。例如，查询“生成答案时外部资料怎样减少胡编”对应文档D25“检索增强生成通过外部知识补充模型参数知识，可缓解知识过时和幻觉问题”；查询“只看字面词会被同义说法误导吗”对应文档D24“传统关键词模型对短查询十分敏感，当用户使用近义表达时可能把真正相关文档排在后面”。这些样例能够体现语义重排序的应用价值。",
    )
    add_image(doc, "case_table.png", "图2  部分查询与相关文档标注截图", width=15.8)
    h(doc, "4.3 对比方法", 2)
    p(
        doc,
        "实验设置两个对比方法。第一种为TF-IDF原始排序，即直接使用字符n-gram TF-IDF和余弦相似度进行排序，不做语义扩展和降维。该方法代表传统词面匹配基线。第二种为TF-IDF+语义重排序，即先保留TF-IDF得到的候选列表，再对查询和文档做同义词归一、LSA语义向量计算和得分融合。两种方法使用相同语料和查询，保证比较公平。",
    )
    h(doc, "4.4 评价指标", 2)
    p(
        doc,
        "Top-1准确率表示第一条结果是否为标准相关文档，适合衡量用户最直接看到的结果质量。Hit@K表示标准相关文档是否出现在前K条结果中，本文报告Hit@3和Hit@5。MRR即平均倒数排名，若相关文档排第1，则该查询得分为1；若排第2，则得分为1/2；若排第10，则得分为1/10。nDCG考虑排名位置折损，相关文档越靠前，得分越高。多个指标结合使用可以避免单一指标片面化。",
    )

    h(doc, "5 实验结果与分析")
    h(doc, "5.1 总体结果", 2)
    metrics = add_metric_table(doc)
    add_image(doc, "metric_table.png", "图3  检索指标对比表截图", width=15.8)
    add_image(doc, "metric_chart.png", "图4  主要指标柱状图", width=14.5)
    p(
        doc,
        "从总体结果看，语义重排序在所有主要指标上均优于原始TF-IDF。Top-1准确率由0.667提升至0.750，说明更多查询能够在第一位获得正确文档；Hit@3由0.833提升至0.917，Hit@5由0.833提升至1.000，说明重排序后标准相关文档更稳定地进入用户更可能浏览的结果范围；MRR由0.742提升至0.840，nDCG@5由0.761提升至0.880，进一步证明相关文档整体排名位置得到改善。",
    )
    p(
        doc,
        "这些提升说明，在中文短查询环境下，单纯依靠字符重合进行排序会遗漏一部分语义相关结果，而同义词归一和潜在语义向量能够提供补充信号。需要注意的是，本文实验规模较小，提升幅度不能直接等同于大规模真实系统效果；但作为方法验证，结果已经能够支持“语义重排序可以改善靠前结果相关性”的基本结论。",
    )
    h(doc, "5.2 查询级结果", 2)
    rank = add_rank_table(doc)
    add_image(doc, "rank_change.png", "图5  查询级排名变化截图", width=15.8)
    p(
        doc,
        "查询级结果能够更直观地展示重排序的作用。以Q08“前K列表过短会怎样限制精排上限”为例，原始TF-IDF把D40排在第2位，而语义重排序后D40上升到第1位。这说明语义模块识别到“前K列表过短”“限制精排上限”与“候选集大小影响重排序上限”之间的对应关系。对于用户而言，相关文档从第二位提升到第一位看似幅度不大，但在问答系统中可能决定生成模型首先读取哪一段上下文。",
    )
    p(
        doc,
        "再看Q12“同时利用字面匹配和向量理解的方案”，标准相关文档为D15。由于D15中包含“混合检索”“稀疏关键词匹配”“稠密向量匹配”等表达，原始TF-IDF已经能够把它排到第一位，重排序保持了该结果。这说明语义重排序并不一定要改变所有查询的排序，合理的重排序方法应当在词面匹配已经可靠时保持稳定，在词面匹配不足时进行纠偏。",
    )
    p(
        doc,
        "也存在重排序未完全解决的问题。例如Q01和Q03中，相关文档在原始排序中位置较靠后，说明初始词面特征已经严重低估其相关性。虽然语义扩展能够提供一定帮助，但由于本文没有使用真正的大规模预训练语义模型，LSA对复杂语义的理解能力有限，无法在所有案例中把相关文档提升到前几位。这一现象也说明，重排序效果受到候选召回质量、语义表示能力和领域词表覆盖率的共同制约。",
    )
    h(doc, "5.3 方法有效原因", 2)
    p(
        doc,
        "本文方法有效的原因首先来自两阶段结构。初始TF-IDF能够快速利用字面线索找到一批候选结果，避免语义模块在全部文档上进行无差别计算；语义重排序则在候选范围内补充近义表达和隐含主题信息。其次，同义词归一把“搜索、查找、检索”“精排、重排序、排序优化”“幻觉、胡编、外部知识”等表达拉近，直接缓解了中文短查询中的词汇不匹配。最后，LSA降维减少了稀疏词项空间中的噪声，使相似主题文档在低维空间中更容易接近。",
    )
    p(
        doc,
        "从工程角度看，该方法还具有较强可解释性。相比端到端神经网络模型，同义词组和加权融合策略更容易被人工检查和调整。当某个查询排序效果不好时，可以追溯是初始召回漏掉、同义词覆盖不足、语义向量表示不充分，还是融合权重设置不合理。这种可解释性对于课程实验和小型知识库项目非常有价值。",
    )
    h(doc, "5.4 局限性分析", 2)
    p(
        doc,
        "本文实验仍存在明显局限。第一，数据集规模较小，只有40条文档和12个查询，不能代表真实搜索系统中的复杂分布。第二，每个查询只标注1条最相关文档，而真实检索任务往往存在多个不同相关等级的文档。第三，同义词归一依赖人工构造，覆盖范围有限，遇到未收录表达时效果会下降。第四，LSA属于较早期语义方法，无法像BERT类模型一样充分利用上下文信息。第五，实验没有测量延迟和资源消耗，因此还不能完整评估工程部署成本。",
    )
    p(
        doc,
        "这些局限并不否定实验结论，而是说明本文更适合作为语义重排序思想的教学复现。若要扩展为更严谨的科研实验，需要使用更大规模公开数据集或真实业务日志，引入多级相关性标注，并与BM25、Sentence-BERT、BGE、Cross-Encoder和大语言模型重排序器等方法进行更全面比较。",
    )

    h(doc, "6 应用讨论")
    h(doc, "6.1 在智能问答系统中的应用", 2)
    p(
        doc,
        "在智能问答系统中，检索模块通常负责为生成模型提供依据。如果检索结果排序较差，生成模型可能读取到无关片段，从而出现答非所问或编造内容。语义重排序可以在生成前对候选上下文进行筛选，把更可能回答用户问题的文档放在前面。对于企业知识库问答，重排序还关系到答案可追溯性，因为用户往往需要知道答案依据来自哪份制度、哪条规范或哪段说明。",
    )
    h(doc, "6.2 在校园和企业知识库中的应用", 2)
    p(
        doc,
        "校园图书馆、课程资料库和企业知识库通常包含大量结构不统一的短文档。用户查询也常常不是标准关键词，而是自然语言问题。例如学生可能输入“怎么找和人工智能检索有关的资料”，员工可能输入“报销制度里交通费怎么处理”。这类查询与文档标题或正文不一定完全重合，语义重排序能够提高系统对意图的理解能力，降低用户反复修改关键词的成本。",
    )
    h(doc, "6.3 与大模型重排序的关系", 2)
    p(
        doc,
        "本文采用的LSA语义重排序属于轻量级方案。若系统拥有更强算力和模型资源，可以把语义模块替换为预训练句向量模型或大语言模型重排序器。句向量模型适合批量计算和向量索引，Cross-Encoder适合小候选集精排，大语言模型则可以通过提示词判断候选文档是否真正回答了问题。它们的共同思想仍然是：先用高效方法召回候选，再用更强语义能力进行精细排序。",
    )

    h(doc, "7 结论与展望")
    p(
        doc,
        "本文围绕信息检索结果优化问题，设计并复现了一种基于语义重排序的两阶段检索方法。该方法首先使用字符n-gram TF-IDF完成初始排序，再通过同义词归一、潜在语义分析和余弦相似度进行候选重排序。实验结果表明，相比原始TF-IDF排序，语义重排序在Top-1准确率、Hit@K、MRR和nDCG等指标上均有所提升，能够使相关文档更靠前地呈现给用户。研究说明，在中文短查询和问答式检索场景中，引入语义信息是提升排序质量的有效途径。",
    )
    p(
        doc,
        "未来可以从四个方向继续改进。第一，扩大数据集规模，使用公开中文检索数据集或真实业务日志进行验证。第二，引入更强的中文句向量模型，如BGE、text2vec或多语言Sentence-BERT，替换本文的LSA表示。第三，在候选集较小的情况下使用Cross-Encoder或大语言模型进行精排，以获得更强的语义交互能力。第四，把重排序模块与RAG系统结合，进一步评估检索排序对最终生成答案准确率、引用可靠性和用户满意度的影响。",
    )

    h(doc, "参考文献")
    refs = [
        "庞亮, 邓竞成, 顾佳, 沈华伟, 程学旗. 大语言模型时代的信息检索综述[A]. 第23届中国计算语言学大会论文集, 2024:98-119. https://aclanthology.org/2024.ccl-2.6/",
        "冯掌印, 朱坤, 马伟涛, 黄磊, 秦兵, 刘挺, 冯骁骋. 浅谈大模型时代下的检索增强:发展趋势、挑战与展望[A]. 第23届中国计算语言学大会论文集, 2024:151-168. https://aclanthology.org/2024.ccl-2.9/",
        "Zhao P, Zhang H, Yu Q, Wang Z, Geng Y, Fu F, Yang L. Retrieval-Augmented Generation for AI-Generated Content: A Survey[J]. Data Science and Engineering, 2026, 11:1-29.",
        "Gao Y, Xiong Y, Gao X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[J/OL]. arXiv:2312.10997, 2023.",
        "Thakur N, Reimers N, Rücklé A, Srivastava A, Gurevych I. BEIR: A Heterogeneous Benchmark for Zero-shot Evaluation of Information Retrieval Models[C]. NeurIPS Datasets and Benchmarks, 2021.",
        "Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]. EMNLP-IJCNLP, 2019.",
        "Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009.",
        "Manning C D, Raghavan P, Schütze H. Introduction to Information Retrieval[M]. Cambridge University Press, 2008.",
        "Karpukhin V, Oguz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]. EMNLP, 2020.",
        "Nogueira R, Cho K. Passage Re-ranking with BERT[J/OL]. arXiv:1901.04085, 2019.",
        "Sentence Transformers Documentation. Semantic Textual Similarity[EB/OL]. https://www.sbert.net/docs/sentence_transformer/usage/semantic_textual_similarity.html.",
    ]
    for ref in refs:
        plain(doc, ref)

    h(doc, "附录A 实验复现说明")
    p(
        doc,
        "实验复现文件包括reproduce_semantic_rerank.py、outputs/documents.csv、outputs/queries.csv、outputs/metrics.csv和outputs/rank_change.csv。运行命令为python reproduce_semantic_rerank.py。脚本会自动生成实验数据、指标表、流程图、柱状图和排名变化截图。扩展版论文生成命令为python generate_expanded_paper.py。由于实验采用本地轻量方法，不需要联网下载预训练模型，因此在普通Python环境下即可复现。",
    )
    p(
        doc,
        "为保证实验可解释性，脚本中直接给出了文档集合、查询集合、同义词归一规则和指标计算方式。读者可以通过修改DOCS、QUERIES和synonym_groups扩展实验语料，也可以调整融合权重观察排序结果变化。若要进一步接近真实系统，可把LSA模块替换为sentence-transformers生成的句向量，再保持相同指标计算流程进行比较。",
    )

    doc.add_page_break()
    h(doc, "附录B 实验复现代码")
    p(
        doc,
        "以下为本文实验复现所使用的Python代码，对应工作目录中的reproduce_semantic_rerank.py。代码包含实验语料构建、TF-IDF初始检索、语义重排序、指标计算、图表生成和论文初版生成过程。",
    )
    code = Path("reproduce_semantic_rerank.py").read_text(encoding="utf-8")
    code_para = doc.add_paragraph()
    code_para.paragraph_format.first_line_indent = Pt(0)
    code_para.paragraph_format.line_spacing = 1.0
    code_run = code_para.add_run(code)
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(7)

    doc.save(PAPER)
    text_len = sum(len(paragraph.text.strip()) for paragraph in doc.paragraphs)
    print(f"Generated: {PAPER}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Text characters: {text_len}")


if __name__ == "__main__":
    build()
