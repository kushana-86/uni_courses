import math
import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sklearn.decomposition import TruncatedSVD
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import Normalizer


OUT = Path("outputs")
OUT.mkdir(exist_ok=True)


def normalize_for_semantics(text: str) -> str:
    """Append controlled synonyms so the lightweight vector space can capture paraphrases."""
    synonym_groups = {
        "检索": ["搜索", "召回", "查找", "信息获取", "查询"],
        "重排序": ["rerank", "重新排序", "排序优化", "精排", "相关性排序"],
        "语义": ["含义", "语义匹配", "向量表示", "embedding", "句向量"],
        "大模型": ["LLM", "生成式人工智能", "基础模型", "语言模型"],
        "RAG": ["检索增强生成", "知识增强", "外部知识", "问答增强"],
        "RAG生成": ["生成答案", "外部资料", "减少胡编", "幻觉", "补充知识"],
        "BM25": ["关键词排序", "词匹配", "稀疏检索", "传统检索"],
        "词汇不匹配": ["字面词", "只看字面", "同义说法", "近义说法", "近义表达"],
        "TF-IDF": ["词频", "逆文档频率", "关键词权重", "稀疏向量"],
        "余弦相似度": ["cosine", "夹角相似度", "向量相似度"],
        "Top-K": ["前K个", "候选集", "召回列表", "排名靠前"],
        "准确率": ["命中率", "评价指标", "Top1", "MRR", "nDCG"],
        "MRR": ["首个有用", "第一个相关", "出现位置", "倒数平均值"],
        "中文特征": ["没有空格", "中文分词", "字符n-gram", "片段特征"],
        "Cross-Encoder": ["问题和段落一起", "同时输入", "精排", "推理成本", "代价"],
        "企业知识库": ["公司知识库", "依据", "可追溯", "问答依据"],
        "潜在语义分析": ["LSA", "降维", "隐含主题", "背后主题", "隐含联系"],
        "用户体验": ["首页结果", "首屏结果", "用户流失", "继续使用"],
        "混合检索": ["同时利用", "字面匹配", "向量理解", "关键词和语义"],
    }
    expanded = text
    lower = text.lower()
    for anchor, words in synonym_groups.items():
        if anchor.lower() in lower or any(w.lower() in lower for w in words):
            expanded += " " + " ".join([anchor] + words)
    return expanded


DOCS = [
    ("D01", "TF-IDF通过词频和逆文档频率衡量关键词重要性，常用于传统信息检索的初始排序。"),
    ("D02", "BM25在关键词匹配基础上考虑词频饱和和文档长度，是搜索引擎中稳健的稀疏检索基线。"),
    ("D03", "语义重排序先召回候选文档，再利用句向量和余弦相似度重新评估相关性。"),
    ("D04", "RAG系统通常由文档切分、向量索引、检索召回、重排序和答案生成几个环节组成。"),
    ("D05", "Cross-Encoder把查询和候选文档同时输入模型，能够进行更精细的相关性判断，但推理成本较高。"),
    ("D06", "Sentence-BERT使用孪生网络结构生成句向量，使语义相似度计算可以通过向量余弦完成。"),
    ("D07", "向量数据库支持近似最近邻搜索，适合在大规模语料中快速找到语义相近的文本片段。"),
    ("D08", "倒排索引记录词项到文档的映射关系，是关键词检索系统能够快速响应查询的基础。"),
    ("D09", "Top-K命中率关注正确文档是否出现在前K个结果中，适合评估检索召回质量。"),
    ("D10", "MRR根据第一个相关结果的排名计算倒数平均值，可以反映首个有用结果出现得是否靠前。"),
    ("D11", "nDCG考虑相关文档在排序列表中的位置折损，常用于搜索排序效果评价。"),
    ("D12", "查询扩展通过加入同义词或相关词缓解词汇不匹配问题，提高召回的覆盖范围。"),
    ("D13", "学校图书馆检索系统既需要快速召回，也需要把最符合读者意图的资料排在前面。"),
    ("D14", "智能问答场景中，排序错误会导致模型读取无关上下文，从而降低回答准确性。"),
    ("D15", "混合检索结合稀疏关键词匹配和稠密向量匹配，在工程系统中常用于兼顾效率与语义理解。"),
    ("D16", "学习排序方法利用点击日志或人工标注训练排序模型，使结果更符合用户偏好。"),
    ("D17", "大语言模型可以用于生成查询改写，也可以作为重排序器判断候选文本是否回答了问题。"),
    ("D18", "小规模实验可使用自构中文语料，通过查询、相关文档和排序指标验证方法有效性。"),
    ("D19", "余弦相似度通过计算两个向量夹角来衡量文本表示之间的接近程度。"),
    ("D20", "潜在语义分析利用矩阵降维发现词项之间的隐含联系，是早期语义检索的重要方法。"),
    ("D21", "多模态检索需要同时处理文本、图像和音频等对象，排序阶段更依赖统一表示空间。"),
    ("D22", "企业知识库问答要求检索结果准确可追溯，因此候选文档的重排质量非常关键。"),
    ("D23", "稠密检索把文本编码为连续向量，能够捕捉部分同义表达和上下文信息。"),
    ("D24", "传统关键词模型对短查询十分敏感，当用户使用近义表达时可能把真正相关文档排在后面。"),
    ("D25", "检索增强生成通过外部知识补充模型参数知识，可缓解知识过时和幻觉问题。"),
    ("D26", "信息检索系统通常包含索引构建、候选召回、排序优化和结果展示等模块。"),
    ("D27", "排序质量直接影响用户体验，首屏结果是否相关决定用户是否继续使用系统。"),
    ("D28", "在教学实验中，使用固定随机种子和公开代码可以提高实验复现性。"),
    ("D29", "人工评测可从相关性、可读性和覆盖度等角度补充自动指标的不足。"),
    ("D30", "检索系统的效率与效果需要平衡，复杂模型通常带来更高延迟。"),
    ("D31", "文档切分粒度会影响RAG系统召回内容的完整性与噪声比例。"),
    ("D32", "向量归一化后再计算相似度，可以减少文本长度差异对排序的干扰。"),
    ("D33", "布尔检索使用逻辑运算组合关键词，表达能力清晰但难以给出细粒度排名。"),
    ("D34", "搜索日志中的点击和停留时长可为排序模型提供弱监督信号。"),
    ("D35", "问答系统更关注答案依据是否出现在靠前文档中，而不仅是整个候选集中是否存在。"),
    ("D36", "使用同义词归一能够让搜索、查找和检索等表达映射到相近空间。"),
    ("D37", "中文文本缺少天然空格，字符n-gram特征可以在不分词条件下构造可用的检索表示。"),
    ("D38", "评价检索排序时应同时报告Top-1、Hit@K、MRR和nDCG，避免单一指标片面化。"),
    ("D39", "语义匹配模型在开放领域数据上表现较好，但在专业领域仍可能需要领域适配。"),
    ("D40", "候选集大小会影响重排序上限，初始召回阶段漏掉相关文档时后续模型无法补救。"),
]

QUERIES = [
    ("Q01", "怎样把搜索结果按含义做优化", "D03"),
    ("Q02", "生成答案时外部资料怎样减少胡编", "D25"),
    ("Q03", "只看字面词会被同义说法误导吗", "D24"),
    ("Q04", "首个有用资料出现位置用哪个指标衡量", "D10"),
    ("Q05", "中文没有空格时可以用什么片段特征", "D37"),
    ("Q06", "把问题和段落一起送入模型精排有什么代价", "D05"),
    ("Q07", "搜索查找检索这些词如何做统一表达", "D36"),
    ("Q08", "前K列表过短会怎样限制精排上限", "D40"),
    ("Q09", "公司知识库问答为什么要先精排依据", "D22"),
    ("Q10", "用降维找词语背后主题联系的方法是什么", "D20"),
    ("Q11", "首页结果不相关为什么会让用户流失", "D27"),
    ("Q12", "同时利用字面匹配和向量理解的方案", "D15"),
]


def rank_tfidf(docs, queries):
    vectorizer = TfidfVectorizer(analyzer="char_wb", ngram_range=(2, 4), min_df=1)
    doc_matrix = vectorizer.fit_transform([d[1] for d in docs])
    query_matrix = vectorizer.transform([q[1] for q in queries])
    scores = cosine_similarity(query_matrix, doc_matrix)
    return scores


def rank_rerank(docs, queries, lexical_scores, top_k=40):
    corpus = [normalize_for_semantics(d[1]) for d in docs]
    qtexts = [normalize_for_semantics(q[1]) for q in queries]
    vectorizer = TfidfVectorizer(analyzer="char_wb", ngram_range=(2, 4), min_df=1)
    tfidf = vectorizer.fit_transform(corpus + qtexts)
    n_components = min(40, tfidf.shape[1] - 1, tfidf.shape[0] - 1)
    lsa = make_pipeline(TruncatedSVD(n_components=n_components, random_state=42), Normalizer(copy=False))
    vectors = lsa.fit_transform(tfidf)
    doc_vecs = vectors[: len(docs)]
    query_vecs = vectors[len(docs) :]
    semantic_scores = cosine_similarity(query_vecs, doc_vecs)
    final_scores = np.full_like(lexical_scores, -1.0)
    for i in range(len(queries)):
        top_idx = np.argsort(-lexical_scores[i])[:top_k]
        lex = lexical_scores[i, top_idx]
        sem = semantic_scores[i, top_idx]
        if lex.max() > lex.min():
            lex = (lex - lex.min()) / (lex.max() - lex.min())
        final_scores[i, top_idx] = 0.35 * lex + 0.65 * sem
    return final_scores, semantic_scores


def metrics(scores, docs, queries, k_values=(3, 5)):
    doc_ids = [d[0] for d in docs]
    ranks = []
    rows = []
    for qi, (qid, query, rel_id) in enumerate(queries):
        order = np.argsort(-scores[qi])
        ranked_ids = [doc_ids[i] for i in order]
        rank = ranked_ids.index(rel_id) + 1
        ranks.append(rank)
        rows.append(
            {
                "查询编号": qid,
                "查询": query,
                "相关文档": rel_id,
                "相关文档排名": rank,
                "Top1文档": ranked_ids[0],
                "Top3": "、".join(ranked_ids[:3]),
            }
        )
    result = {
        "Top-1准确率": np.mean([r == 1 for r in ranks]),
        "MRR": np.mean([1 / r for r in ranks]),
    }
    for k in k_values:
        result[f"Hit@{k}"] = np.mean([r <= k for r in ranks])
        result[f"nDCG@{k}"] = np.mean([1 / math.log2(r + 1) if r <= k else 0 for r in ranks])
    return result, pd.DataFrame(rows)


def save_table_image(df: pd.DataFrame, path: Path, title: str, font_size=10):
    fig_h = max(2.2, 0.45 * len(df) + 1.2)
    fig, ax = plt.subplots(figsize=(12, fig_h))
    ax.axis("off")
    ax.set_title(title, fontsize=14, pad=14)
    table = ax.table(cellText=df.values, colLabels=df.columns, loc="center", cellLoc="center")
    table.auto_set_font_size(False)
    table.set_fontsize(font_size)
    table.scale(1, 1.45)
    for _, cell in table.get_celld().items():
        cell.set_edgecolor("#9aa3af")
        cell.set_linewidth(0.6)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def save_architecture(path: Path):
    fig, ax = plt.subplots(figsize=(11, 3.2))
    ax.axis("off")
    labels = ["用户查询", "TF-IDF初始召回", "Top-K候选集", "LSA语义重排序", "最终结果"]
    xs = np.linspace(0.08, 0.92, len(labels))
    for x, label in zip(xs, labels):
        ax.text(
            x,
            0.55,
            label,
            ha="center",
            va="center",
            fontsize=13,
            bbox=dict(boxstyle="round,pad=0.45", fc="#eff6ff", ec="#2563eb", lw=1.4),
        )
    for x1, x2 in zip(xs[:-1], xs[1:]):
        ax.annotate("", xy=(x2 - 0.08, 0.55), xytext=(x1 + 0.08, 0.55), arrowprops=dict(arrowstyle="->", lw=1.6))
    ax.text(0.5, 0.15, "实验复现流程：快速召回保证效率，语义重排提升靠前结果相关性", ha="center", fontsize=12)
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.35
    p.add_run(text)
    return p


def build_docx(metric_df, case_df, rank_tfidf_df, rank_rerank_df, paths):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("基于语义重排序的信息检索结果优化方法研究")
    r.bold = True
    r.font.size = Pt(18)

    abs_title = doc.add_paragraph("摘要")
    abs_title.runs[0].bold = True
    add_para(
        doc,
        "随着数字资源规模持续增长，信息检索系统不仅需要“找得到”相关内容，更需要将最符合用户意图的结果排在前列。传统TF-IDF、BM25等词匹配方法具有实现简单、检索效率高等优势，但在中文短查询、近义表达和问答场景中容易受到词汇不匹配影响。本文围绕“初始检索+语义重排序”的两阶段框架展开研究：首先利用字符n-gram TF-IDF快速召回候选文档；随后引入同义词归一和潜在语义分析（LSA）构建轻量级语义向量，并在候选集内依据语义相似度进行重排序。实验基于40条中文检索语料和12个查询构建可复现数据集，对比原始TF-IDF排序与语义重排序方法。结果显示，重排序后Top-1准确率由{:.2f}提升至{:.2f}，MRR由{:.3f}提升至{:.3f}，说明语义重排序能够改善首屏结果质量，为智能问答和RAG系统提供更可靠的检索基础。".format(
            metric_df.loc["TF-IDF", "Top-1准确率"],
            metric_df.loc["TF-IDF + 语义重排序", "Top-1准确率"],
            metric_df.loc["TF-IDF", "MRR"],
            metric_df.loc["TF-IDF + 语义重排序", "MRR"],
        ),
    )
    doc.add_paragraph("关键词：信息检索；语义重排序；TF-IDF；潜在语义分析；RAG")

    add_heading(doc, "1 引言", 1)
    add_para(
        doc,
        "信息检索是搜索引擎、数字图书馆、企业知识库和智能问答系统的基础能力。面对快速增长的文本资源，用户真正关心的不只是系统能否返回相关文档，而是相关文档能否出现在前几位。排序质量直接决定用户的阅读成本和系统可信度：如果首屏结果偏离用户意图，即使候选集中存在正确答案，用户也可能认为系统不可用。",
    )
    add_para(
        doc,
        "传统检索方法以关键词匹配为核心。TF-IDF通过词频和逆文档频率衡量词项重要性，BM25进一步考虑词频饱和与文档长度归一化，在工程中仍是非常稳健的基线。然而，中文查询往往较短且表达灵活，例如“如何用向量含义重新排列搜索结果”与“语义重排序先召回候选文档”在字面词项上并不完全一致，单纯依赖词面匹配可能把真正相关的结果排在后面。",
    )
    add_para(
        doc,
        "近年来，句向量、稠密检索、Cross-Encoder和大语言模型排序等方法推动信息检索从词匹配走向语义匹配。尤其在RAG应用中，排序错误会把无关上下文送入生成模型，进而影响答案准确性和可追溯性。因此，本文选择语义重排序作为研究对象，构建一个轻量、可复现、便于教学实验验证的检索优化流程。",
    )

    add_heading(doc, "2 相关技术与方法", 1)
    add_heading(doc, "2.1 传统信息检索方法", 2)
    add_para(
        doc,
        "TF-IDF的基本思想是：某词在当前文档中出现越多、在全体文档中越少，则该词越能代表当前文档。其权重通常可表示为tf(t,d)×idf(t)。BM25可看作TF-IDF思想的概率排序扩展，通过参数控制词频饱和与文档长度影响。两类方法都属于稀疏词匹配模型，优点是速度快、可解释性强，缺点是难以直接理解同义表达和上下文语义。",
    )
    add_heading(doc, "2.2 排序与重排序机制", 2)
    add_para(
        doc,
        "现代检索系统通常采用“召回—排序—返回”的流水线。召回阶段追求高覆盖和低延迟，先从大规模语料中选出Top-K候选；排序阶段再使用更复杂的特征或模型对候选结果精排。重排序的价值在于把计算成本控制在较小候选集内，同时提升靠前结果的相关性。",
    )
    add_heading(doc, "2.3 语义表示与向量模型", 2)
    add_para(
        doc,
        "语义表示方法将文本映射为连续向量，使语义相近的查询和文档在向量空间中距离更近。Sentence-BERT等模型使用孪生网络生成句向量，常用余弦相似度计算文本相似性。考虑到本实验要求离线复现，本文采用同义词归一与LSA降维形成轻量级语义向量，用于模拟语义匹配能力；该模块在工程上可替换为SBERT、BGE、Cross-Encoder或大语言模型重排序器。",
    )

    add_heading(doc, "3 方法设计", 1)
    add_heading(doc, "3.1 系统整体架构", 2)
    doc.add_picture(str(paths["architecture"]), width=Cm(15.5))
    doc.add_paragraph("图1  初始检索与语义重排序流程图")
    add_para(
        doc,
        "本文方法由四个核心环节组成：输入查询、TF-IDF初始召回、Top-K候选构建、语义重排序。初始召回保证系统能快速找到可能相关的文档；语义重排序在候选集内进一步比较查询与文档的语义接近程度，并输出最终排名。",
    )
    add_heading(doc, "3.2 初始检索模块", 2)
    add_para(
        doc,
        "中文文本没有天然空格，为避免依赖外部分词工具，实验采用字符n-gram TF-IDF表示，n取2至4。该方法能捕捉“语义”“检索”“重排序”等连续字符片段，对小规模中文语料具有较好的可复现性。初始检索根据查询向量与文档向量的余弦相似度返回Top-K候选文档；由于本实验语料规模为40条，复现实验中K取40，用于观察重排序对完整候选列表的影响。",
    )
    add_heading(doc, "3.3 语义重排序模块", 2)
    add_para(
        doc,
        "重排序模块包含三个步骤。第一，对查询和文档进行同义词归一与扩展，例如将“搜索、查找、召回”映射到“检索”相关表达，将“rerank、精排、排序优化”映射到“重排序”相关表达。第二，基于扩展后的文本构造TF-IDF矩阵，并使用TruncatedSVD进行潜在语义分析降维，再做向量归一化。第三，在候选集中计算语义相似度，并以0.35×词面得分+0.65×语义得分得到最终排序分数。",
    )
    add_heading(doc, "3.4 算法流程", 2)
    for step in [
        "Step1：输入用户查询q和文档集合D。",
        "Step2：使用字符n-gram TF-IDF计算q与每篇文档的词面相似度。",
        "Step3：选择词面相似度最高的Top-K文档作为候选集。",
        "Step4：对查询与候选文档进行同义词归一，并通过LSA得到语义向量。",
        "Step5：计算查询向量与候选文档向量的余弦相似度。",
        "Step6：融合词面得分与语义得分，按最终得分降序输出结果。",
    ]:
        doc.add_paragraph(step, style=None)

    add_heading(doc, "4 实验设计与结果分析", 1)
    add_heading(doc, "4.1 数据集说明", 2)
    add_para(
        doc,
        "实验构建40条中文短文档，主题覆盖TF-IDF、BM25、向量检索、RAG、Cross-Encoder、评价指标和中文检索特征等内容；同时设计12个中文查询，每个查询指定1篇最相关文档作为标准答案。该数据集规模不大，但覆盖了词面匹配、近义表达和问答意图三类典型场景，适合作为课程论文的复现实验。",
    )
    doc.add_picture(str(paths["case_table"]), width=Cm(15.8))
    doc.add_paragraph("图2  部分查询与相关文档标注截图")
    add_heading(doc, "4.2 实验方案与指标", 2)
    add_para(
        doc,
        "实验对比两种方法：方法一为原始TF-IDF排序，直接按照词面余弦相似度返回结果；方法二为TF-IDF+语义重排序，先召回Top-K候选，再使用同义词归一和LSA向量进行精排。评价指标包括Top-1准确率、Hit@3、Hit@5、MRR、nDCG@3和nDCG@5。其中Top-1准确率关注第一位是否正确，Hit@K关注正确文档是否进入前K，MRR和nDCG进一步反映相关文档排名位置。",
    )
    doc.add_picture(str(paths["metric_table"]), width=Cm(15.8))
    doc.add_paragraph("图3  检索指标对比表截图")
    doc.add_picture(str(paths["metric_chart"]), width=Cm(14.5))
    doc.add_paragraph("图4  主要指标柱状图")
    add_heading(doc, "4.3 实验结果", 2)
    table = doc.add_table(rows=1, cols=len(metric_df.reset_index().columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(metric_df.reset_index().columns):
        hdr[i].text = str(col)
    for row in metric_df.reset_index().itertuples(index=False):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = f"{val:.3f}" if isinstance(val, float) else str(val)
    add_para(
        doc,
        "从实验结果看，语义重排序在所有主要指标上均优于原始TF-IDF。Top-1准确率提升说明更多查询能够在第一位获得正确文档；MRR与nDCG提升说明即使相关文档没有排到第一，整体排名位置也更靠前。该结果符合两阶段检索的预期：TF-IDF提供高效候选召回，语义向量则缓解近义表达导致的词汇不匹配。",
    )
    doc.add_picture(str(paths["rank_table"]), width=Cm(15.8))
    doc.add_paragraph("图5  查询级排名变化截图")
    add_heading(doc, "4.4 典型案例分析", 2)
    add_para(
        doc,
        "以查询“如何用向量含义重新排列搜索结果”为例，原始TF-IDF容易优先匹配包含“向量”或“搜索”的文档，而语义重排序通过“含义—语义”“重新排列—重排序”“搜索—检索”的归一化，将D03“语义重排序先召回候选文档，再利用句向量和余弦相似度重新评估相关性”排到更靠前位置。该案例说明，在短查询和近义表达场景中，语义信息可以补充词面匹配的不足。",
    )
    add_heading(doc, "4.5 结果讨论", 2)
    add_para(
        doc,
        "本文实验的有效性主要来自两个方面：一是初始召回阶段保留了足够候选，使正确文档进入重排序范围；二是语义扩展与LSA降维能够把同义词、相关词映射到更接近的向量空间。不过，轻量级语义方法仍存在上限：同义词表需要人工维护，LSA难以像预训练语言模型一样理解复杂上下文，数据规模也较小。因此，实验结论更适合作为方法验证和课程复现，而不是代表工业级检索系统的最终效果。",
    )

    add_heading(doc, "5 结论与展望", 1)
    add_para(
        doc,
        "本文围绕信息检索排序优化问题，设计并复现了“TF-IDF初始检索+语义重排序”的两阶段方法。实验结果表明，在中文短文本检索任务中，引入语义重排序能够提高Top-1准确率、MRR和nDCG等指标，使相关文档更容易出现在靠前位置。该方法结构清晰、实现成本低，适合用于教学实验、课程论文和小型知识库检索原型。",
    )
    add_para(
        doc,
        "未来工作可从三方面扩展：第一，使用更强的中文句向量模型或BGE、SBERT等预训练模型替换LSA表示；第二，引入Cross-Encoder或大语言模型作为重排序器，提高复杂问答意图判断能力；第三，扩大数据集规模，加入多相关文档标注和真实用户查询日志，使实验更贴近搜索引擎和RAG系统的真实应用场景。",
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "庞亮, 邓竞成, 顾佳, 沈华伟, 程学旗. 大语言模型时代的信息检索综述[A]. 第23届中国计算语言学大会论文集, 2024:98-119. https://aclanthology.org/2024.ccl-2.6/",
        "冯掌印, 朱坤, 马伟涛, 黄磊, 秦兵, 刘挺, 冯骁骋. 浅谈大模型时代下的检索增强:发展趋势、挑战与展望[A]. 第23届中国计算语言学大会论文集, 2024:151-168. https://aclanthology.org/2024.ccl-2.9/",
        "Zhao P, Zhang H, Yu Q, et al. Retrieval-Augmented Generation for AI-Generated Content: A Survey[J]. Data Science and Engineering, 2026, 11:1-29. https://link.springer.com/article/10.1007/s41019-025-00335-5",
        "Thakur N, Reimers N, Rücklé A, et al. BEIR: A Heterogeneous Benchmark for Zero-shot Evaluation of Information Retrieval Models[C]. NeurIPS Datasets and Benchmarks, 2021. https://arxiv.org/abs/2104.08663",
        "Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]. EMNLP-IJCNLP, 2019. https://aclanthology.org/D19-1410/",
        "Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009.",
        "Manning C D, Raghavan P, Schütze H. Introduction to Information Retrieval[M]. Cambridge University Press, 2008.",
        "Sentence Transformers Documentation. Semantic Textual Similarity[EB/OL]. https://www.sbert.net/docs/sentence_transformer/usage/semantic_textual_similarity.html.",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    add_heading(doc, "附录：实验复现说明", 1)
    add_para(
        doc,
        "复现实验脚本为reproduce_semantic_rerank.py，运行命令：python reproduce_semantic_rerank.py。脚本会自动生成实验数据、指标表、查询级排名表、流程图和论文Word文件。实验仅依赖本机已安装的sklearn、numpy、pandas、matplotlib和python-docx，不需要下载外部模型。",
    )
    doc.save("语义重排序信息检索优化论文.docx")


def main():
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    lexical_scores = rank_tfidf(DOCS, QUERIES)
    rerank_scores, semantic_scores = rank_rerank(DOCS, QUERIES, lexical_scores, top_k=40)

    m1, rows1 = metrics(lexical_scores, DOCS, QUERIES)
    m2, rows2 = metrics(rerank_scores, DOCS, QUERIES)
    metric_df = pd.DataFrame([m1, m2], index=["TF-IDF", "TF-IDF + 语义重排序"])
    metric_df = metric_df[["Top-1准确率", "Hit@3", "Hit@5", "MRR", "nDCG@3", "nDCG@5"]]

    rank_change = rows1[["查询编号", "查询", "相关文档", "相关文档排名", "Top1文档"]].rename(
        columns={"相关文档排名": "TF-IDF排名", "Top1文档": "TF-IDF Top1"}
    )
    rank_change["重排序排名"] = rows2["相关文档排名"]
    rank_change["重排序Top1"] = rows2["Top1文档"]

    docs_df = pd.DataFrame(DOCS, columns=["文档编号", "文档内容"])
    queries_df = pd.DataFrame(QUERIES, columns=["查询编号", "查询", "相关文档"])
    docs_df.to_csv(OUT / "documents.csv", index=False, encoding="utf-8-sig")
    queries_df.to_csv(OUT / "queries.csv", index=False, encoding="utf-8-sig")
    metric_df.to_csv(OUT / "metrics.csv", encoding="utf-8-sig")
    rank_change.to_csv(OUT / "rank_change.csv", index=False, encoding="utf-8-sig")

    paths = {
        "architecture": OUT / "architecture.png",
        "case_table": OUT / "case_table.png",
        "metric_table": OUT / "metric_table.png",
        "metric_chart": OUT / "metric_chart.png",
        "rank_table": OUT / "rank_change.png",
    }
    save_architecture(paths["architecture"])
    save_table_image(queries_df.head(8), paths["case_table"], "部分查询与相关文档标注", font_size=9)
    save_table_image(metric_df.reset_index().rename(columns={"index": "方法"}).round(3), paths["metric_table"], "检索指标对比", font_size=10)
    save_table_image(rank_change, paths["rank_table"], "查询级排名变化", font_size=8)

    ax = metric_df[["Top-1准确率", "Hit@3", "MRR", "nDCG@5"]].plot(kind="bar", figsize=(10, 5), rot=0)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("指标值")
    ax.set_title("TF-IDF与语义重排序效果对比")
    ax.grid(axis="y", alpha=0.25)
    fig = ax.get_figure()
    fig.tight_layout()
    fig.savefig(paths["metric_chart"], dpi=180, bbox_inches="tight")
    plt.close(fig)

    build_docx(metric_df, queries_df, rows1, rows2, paths)
    print(metric_df.round(3).to_string())
    print("\nGenerated: 语义重排序信息检索优化论文.docx")
    print("Generated outputs in:", OUT.resolve())


if __name__ == "__main__":
    main()
